import io
from datetime import datetime, timedelta
from api.models import *
from api.serializer import *
from django.conf import settings
from rest_framework.response import Response
from django.core.files.storage import default_storage

# Document Manipulation
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement



# Base url
def base_url(value):
    scheme = value.scheme
    host = value.get_host()
    return f"{scheme}://{host}"


def get_school_folder(school_name: str):
    name = school_name.replace(".", "_").replace(" ", "_").replace("'", "").lower()
    return name


def get_student_transcript(student, request):
    student_data = StudentSerializer(student).data
    student_class = ClasseWithSubjectsSerializer(Classe.objects.get(students=student)).data
    st_name = f"{student_data['user']['first_name']} {student_data['user']['last_name']}"
    st_date_enrolled = datetime.fromisoformat(student_data['date_enrolled']).strftime("%d %B, %Y")
    st_program = student_data['program']['name']
    st_gender = student_data['gender']
    st_id = student_data['st_id']

    if student_data['has_completed']:
        st_graduation_date = datetime.fromisoformat(student_class['completion_date']).strftime("%d %B, %Y")
    else:
        if student_class['completion_date']:
            st_graduation_date = datetime.fromisoformat(student_class['completion_date']).strftime("%B %Y")
        else:
            st_graduation_date = "NOT YET"

    doc = Document()
    section = doc.sections[0]
    section.left_margin = 0
    section.right_margin = 0
    section.top_margin = 0
    section.bottom_margin = 0

    # Font Family Styles
    arial = doc.styles.add_style('arial', 1)  # Arial
    arial.font.name = 'Arial'
    verdana1 = doc.styles.add_style('verdana 1', 1)  # Verdana Paragraph style
    verdana1.font.name = 'Verdana'
    verdana2 = doc.styles.add_style('verdana 2', 2)  # Verdana Run style
    verdana2.font.name = 'Verdana'

    # Top Bar
    top_bar_paragraph = doc.add_paragraph("Powered By: EduAAP")
    top_bar_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    top_bar_paragraph.paragraph_format.space_after = Pt(0)
    top_bar_paragraph.runs[0].font.size = Pt(7.5)
    top_bar_paragraph.runs[0].font.color.rgb = RGBColor(255, 255, 0)
    
    # Top Bar Shade
    top_bar_shd = OxmlElement('w:shd')
    top_bar_shd.set(qn('w:val'), 'clear')
    top_bar_shd.set(qn('w:color'), 'auto')
    top_bar_shd.set(qn('w:fill'), '2E8B57')
    top_bar_paragraph.paragraph_format.element.get_or_add_pPr().append(top_bar_shd)

    # School logo
    sch_logo = doc.add_paragraph()
    sch_logo.add_run().add_picture(f"{student_data['school']['sch_logo'].replace(f'{base_url(request)}/', '')}", height=Inches(1.0))
    sch_logo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    sch_logo.paragraph_format.space_before = Pt(20)
    sch_logo.paragraph_format.space_after = Pt(0)
    # School Logo Shade
    logo_shd = OxmlElement('w:shd')
    logo_shd.set(qn('w:val'), 'clear')
    logo_shd.set(qn('w:color'), 'auto')
    logo_shd.set(qn('w:fill'), 'FFF6D3')
    sch_logo.paragraph_format.element.get_or_add_pPr().append(logo_shd)

    sch_name = doc.add_paragraph(f"{student_data['school']['name']}")
    sch_address = doc.add_paragraph(student_data['school']['address'])
    accreditation = doc.add_paragraph("SCHOOL ACCREDITED BY THE GOVERNMENT")
    trans_date = datetime.fromisoformat(datetime.now().date().isoformat()).strftime("%d %B, %Y")
    sch_info = doc.add_paragraph(f"{trans_date.upper()}")
    trans_title = doc.add_paragraph(f"ACADEMIC TRANSCRIPT TO {st_name.upper()}")

    if student_data['has_completed']:
        graduation_date = f"DATE COMPLETED: {st_graduation_date.upper()}"
    else:
        graduation_date = f"EXPECTED COMPLETION DATE: {st_graduation_date.upper()}"

    date_enrolled = doc.add_paragraph(f"DATE ENROLLED: {st_date_enrolled.upper()}    |    {graduation_date}")
    program = doc.add_paragraph(f"PROGRAM: {st_program.upper()}")
    gender = f"GENDER: {st_gender.upper()}"
    student_id = doc.add_paragraph(f"STUDENT ID: {st_id}    |    {gender}")
    phone = doc.add_paragraph("|")
    # Divider
    divider_para = doc.add_paragraph("divider")
    divider_para.paragraph_format.space_after = Pt(0)

    paragraph_items = [sch_name, sch_address, accreditation, sch_info, trans_title, date_enrolled, program, student_id, phone, divider_para]
    for item in paragraph_items:
        item.style = verdana1
        item.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        item.paragraph_format.space_after = Pt(5)
        item.runs[0].font.size = Pt(8)
        item.runs[0].font.color.rgb = RGBColor(46, 139, 87)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'FFF6D3')
        item.paragraph_format.element.get_or_add_pPr().append(shd)
        if paragraph_items.index(item) == 0:
            item.style = arial
            item.runs[0].font.size = Pt(20)
            item.runs[0].font.bold = True
            item.paragraph_format.space_after = Pt(0)
        elif paragraph_items.index(item) == 4:
            item.style = arial
            item.runs[0].font.size = Pt(10)
            item.runs[0].bold = True

        elif paragraph_items.index(item) == 9:
            item.runs[0].font.size = Pt(10)
            item.paragraph_format.space_after = Pt(0)
            item.runs[0].font.color.rgb = RGBColor(255, 246, 211)

    # Class Record Table
    container = doc.add_table(1, 3)
    for row in container.rows:
        for cell in row.cells:
            cell_xml_element = cell._tc
            cell_shd = OxmlElement('w:shd')
            cell_shd.set(qn('w:val'), 'clear')
            cell_shd.set(qn('w:color'), 'auto')
            cell_shd.set(qn('w:fill'), 'FFF6D3')
            cell_xml_element.get_or_add_tcPr().append(cell_shd)

    # Determine number of semesters in each academic year
    academic_year_1 = None
    academic_year_2 = None
    academic_year_3 = None
    academic_year_names = [year['name'] for year in
                           sorted(student_class['academic_years'], key=lambda x: x['start_date'])]

    academic_year_1 = AcademicYearSerializer(
        AcademicYear.objects.get(school=student.school, name=academic_year_names[0])).data

    if academic_year_names and len(academic_year_names) > 1:
        academic_year_2 = AcademicYearSerializer(
            AcademicYear.objects.get(school=student.school, name=academic_year_names[1])).data

    if academic_year_names and len(academic_year_names) > 2:
        academic_year_3 = AcademicYearSerializer(
            AcademicYear.objects.get(school=student.school, name=academic_year_names[2])).data

    # Inner Table
    table_container_cell = container.rows[0].cells[1]
    container.columns[1].cells[0].width = Inches(13)
    table_container_cell.paragraphs[0].add_run("CLASS RECORD")
    table_container_cell.paragraphs[0].runs[0].font.size = Pt(9)
    table_container_cell.paragraphs[0].runs[0].font.bold = True
    table_container_cell.paragraphs[0].style = arial
    table_container_cell.paragraphs[0].paragraph_format.space_after = Pt(5)
    table_container_cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    table_container_cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(46, 139, 87)
    table_container_cell.paragraphs[0].runs[0].font.underline = True
    r_table = table_container_cell.add_table(rows=2, cols=10)

    for cell in r_table.columns[0].cells:
        cell.width = Inches(5)

    r_table.style = 'TableGrid'
    year = r_table.rows[0].cells[0]
    year.paragraphs[0].add_run('YEAR', style=verdana2)

    if academic_year_1['term_3_start_date']:
        a = r_table.cell(0, 1)
        b = r_table.cell(0, 2)
        c = r_table.cell(0, 3)
        year_1 = a.merge(b).merge(c)
        year_one_sem = r_table.rows[1].cells[1:4]
        if academic_year_2:
            if academic_year_2['term_3_start_date']:
                d = r_table.cell(0, 4)
                e = r_table.cell(0, 5)
                f = r_table.cell(0, 6)
                year_2 = d.merge(e).merge(f)
                year_two_sem = r_table.rows[1].cells[4:7]
                if academic_year_3:
                    if academic_year_3['term_3_start_date']:
                        g = r_table.cell(0, 7)
                        h = r_table.cell(0, 8)
                        i = r_table.cell(0, 9)
                        year_3 = g.merge(h).merge(i)
                        year_three_sem = r_table.rows[1].cells[7:10]
                    else:
                        g = r_table.cell(0, 7)
                        h = r_table.cell(0, 8)
                        year_3 = g.merge(h)
                        year_three_sem = r_table.rows[1].cells[7:9]
                else:
                    g = r_table.cell(0, 7)
                    h = r_table.cell(0, 8)
                    i = r_table.cell(0, 9)
                    year_3 = g.merge(h).merge(i)
                    year_three_sem = r_table.rows[1].cells[7:10]
            else:
                d = r_table.cell(0, 4)
                e = r_table.cell(0, 5)
                year_2 = d.merge(e)
                year_two_sem = r_table.rows[1].cells[4:6]
                if academic_year_3:
                    if academic_year_3['term_3_start_date']:
                        f = r_table.cell(0, 6)
                        g = r_table.cell(0, 7)
                        h = r_table.cell(0, 8)
                        year_3 = f.merge(g).merge(h)
                        year_three_sem = r_table.rows[1].cells[6:9]
                    else:
                        f = r_table.cell(0, 6)
                        g = r_table.cell(0, 7)
                        year_3 = f.merge(g)
                        year_three_sem = r_table.rows[1].cells[6:8]
                else:
                    f = r_table.cell(0, 6)
                    g = r_table.cell(0, 7)
                    year_3 = f.merge(g)
                    year_three_sem = r_table.rows[1].cells[6:8]
        else:
            d = r_table.cell(0, 4)
            e = r_table.cell(0, 5)
            f = r_table.cell(0, 6)
            g = r_table.cell(0, 7)
            h = r_table.cell(0, 8)
            i = r_table.cell(0, 9)
            year_2 = d.merge(e).merge(f)
            year_3 = g.merge(h).merge(i)
            year_two_sem = r_table.rows[1].cells[4:7]
            year_three_sem = r_table.rows[1].cells[7:10]

    else:
        a = r_table.cell(0, 1)
        b = r_table.cell(0, 2)
        year_1 = a.merge(b)
        year_one_sem = r_table.rows[1].cells[1:3]
        if academic_year_2:
            if academic_year_2['term_3_start_date']:
                c = r_table.cell(0, 3)
                d = r_table.cell(0, 4)
                e = r_table.cell(0, 5)
                year_2 = c.merge(d).merge(e)
                year_two_sem = r_table.rows[1].cells[3:6]
                if academic_year_3:
                    if academic_year_3['term_3_start_date']:
                        f = r_table.cell(0, 6)
                        g = r_table.cell(0, 7)
                        h = r_table.cell(0, 8)
                        year_3 = f.merge(g).merge(h)
                        year_three_sem = r_table.rows[1].cells[6:9]
                    else:
                        f = r_table.cell(0, 6)
                        g = r_table.cell(0, 7)
                        year_3 = f.merge(g)
                        year_three_sem = r_table.rows[1].cells[6:8]
                else:
                    f = r_table.cell(0, 6)
                    g = r_table.cell(0, 7)
                    h = r_table.cell(0, 8)
                    year_3 = f.merge(g).merge(h)
                    year_three_sem = r_table.rows[1].cells[6:9]
            else:
                c = r_table.cell(0, 3)
                d = r_table.cell(0, 4)
                year_2 = c.merge(d)
                year_two_sem = r_table.rows[1].cells[3:5]
                if academic_year_3:
                    if academic_year_3['term_3_start_date']:
                        e = r_table.cell(0, 5)
                        f = r_table.cell(0, 6)
                        g = r_table.cell(0, 7)
                        year_3 = e.merge(f).merge(g)
                        year_three_sem = r_table.rows[1].cells[5:8]
                    else:
                        e = r_table.cell(0, 5)
                        f = r_table.cell(0, 6)
                        year_3 = e.merge(f)
                        year_three_sem = r_table.rows[1].cells[5:7]

                        # empty_col_1 = [row.cells[-1] for row in r_table.rows]
                        # empty_col_2 = [row.cells[-2] for row in r_table.rows]
                        # empty_col_3 = [row.cells[-3] for row in r_table.rows]
                        # for cell in empty_col_1:
                        #     cell._element.getparent().remove(cell._element)
                        # for cell in empty_col_2:
                        #     cell._element.getparent().remove(cell._element)
                        # for cell in empty_col_3:
                        #     cell._element.getparent().remove(cell._element)

                else:
                    e = r_table.cell(0, 5)
                    f = r_table.cell(0, 6)
                    year_3 = e.merge(f)
                    year_three_sem = r_table.rows[1].cells[5:7]
        else:
            c = r_table.cell(0, 3)
            d = r_table.cell(0, 4)
            e = r_table.cell(0, 5)
            f = r_table.cell(0, 6)
            year_2 = c.merge(d)
            year_3 = e.merge(f)
            year_two_sem = r_table.rows[1].cells[3:5]
            year_three_sem = r_table.rows[1].cells[5:7]

    year_1.paragraphs[0].add_run('YEAR 1', style=verdana2)
    year_2.paragraphs[0].add_run('YEAR 2', style=verdana2)
    year_3.paragraphs[0].add_run('YEAR 3', style=verdana2)
    subject_term = r_table.rows[1].cells[0]
    subject_term.paragraphs[0].add_run('SUBJECT/SEMESTER', style=verdana2)
    table_head_cell_items = [year, year_1, year_2, year_3, subject_term]

    for cell in table_head_cell_items:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(46, 139, 87)
        cell.paragraphs[0].runs[0].font.size = Pt(8)
        cell.paragraphs[0].runs[0].font.bold = True

    term_count = 1
    for cell1, cell2, cell3 in zip(year_one_sem, year_two_sem, year_three_sem):
        # Term One
        cell1.paragraphs[0].add_run(f"SEMESTER {term_count}", style=verdana2)
        cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell1.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell1.paragraphs[0].runs[0].font.size = Pt(7)
        cell1.paragraphs[0].runs[0].font.bold = True
        cell1.paragraphs[0].runs[0].font.color.rgb = RGBColor(46, 139, 87)

        # Term One
        cell2.paragraphs[0].add_run(f"SEMESTER {term_count}", style=verdana2)
        cell2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell2.paragraphs[0].runs[0].font.color.rgb = RGBColor(46, 139, 87)
        cell2.paragraphs[0].runs[0].font.size = Pt(7)
        cell2.paragraphs[0].runs[0].font.bold = True

        # Term One
        if cell3 is not None:
            cell3.paragraphs[0].add_run(f"SEMESTER {term_count}", style=verdana2)
            cell3.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell3.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell3.paragraphs[0].runs[0].font.color.rgb = RGBColor(46, 139, 87)
            cell3.paragraphs[0].runs[0].font.size = Pt(7)
            cell3.paragraphs[0].runs[0].font.bold = True

        term_count += 1

    # subjects Rows
    for subject in student_class['subjects']:
        new_row = r_table.add_row().cells

        # if academic_year_1['term_3_start_date']:
        #     if academic_year_2:
        #         if academic_year_2['term_3_start_date']:
        #             if academic_year_3:
        #                 if not academic_year_3['term_3_start_date']:
        #                     new_row[-1]._element.getparent().remove(new_row[-1]._element)
        #
        #         else:
        #             if academic_year_3:
        #                 if academic_year_3['term_3_start_date']:
        #                     new_row[-1]._element.getparent().remove(new_row[-1]._element)
        #                 else:
        #                     new_row[-1]._element.getparent().remove(new_row[-1]._element)
        #                     new_row[-1]._element.getparent().remove(new_row[-1]._element)
        #             else:
        #                 new_row[-1]._element.getparent().remove(new_row[-1]._element)
        #                 new_row[-1]._element.getparent().remove(new_row[-1]._element)
        #
        # else:
        #     if academic_year_2:
        #         if academic_year_2['term_3_start_date']:
        #             if academic_year_3:
        #                 if academic_year_3['term_3_start_date']:
        #                     new_row[-1]._element.getparent().remove(new_row[-1]._element)
        #                 else:
        #                     new_row[-1]._element.getparent().remove(new_row[-1]._element)
        #                     new_row[-1]._element.getparent().remove(new_row[-1]._element)
        #             else:
        #                 new_row[-1]._element.getparent().remove(new_row[-1]._element)
        #                 new_row[-1]._element.getparent().remove(new_row[-1]._element)
        #         else:
        #             if academic_year_3:
        #                 if academic_year_3['term_3_start_date']:
        #                     new_row[-1]._element.getparent().remove(new_row[-1]._element)
        #                     new_row[-1]._element.getparent().remove(new_row[-1]._element)
        #                 else:
        #                     new_row[-1]._element.getparent().remove(new_row[-1]._element)
        #                     new_row[-1]._element.getparent().remove(new_row[-1]._element)
        #                     new_row[-1]._element.getparent().remove(new_row[-1]._element)
        #
        #             else:
        #                 new_row[-1]._element.getparent().remove(new_row[-1]._element)
        #                 new_row[-1]._element.getparent().remove(new_row[-1]._element)
        #                 new_row[-1]._element.getparent().remove(new_row[-1]._element)
        #     else:
        #         new_row[-1]._element.getparent().remove(new_row[-1]._element)
        #         new_row[-1]._element.getparent().remove(new_row[-1]._element)
        #         new_row[-1]._element.getparent().remove(new_row[-1]._element)

        new_row[0].paragraphs[0].add_run(f"{subject['name'].upper()}", style=verdana2)
        new_row[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        new_row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        new_row[0].paragraphs[0].runs[0].font.size = Pt(7)
        new_row[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(46, 139, 87)

        if academic_year_1['term_3_start_date']:
            year_one = new_row[1:4]
            if academic_year_2:
                if academic_year_2['term_3_start_date']:
                    year_two = new_row[4:7]
                    if academic_year_3:
                        if academic_year_3['term_3_start_date']:
                            year_three = new_row[7:10]
                        else:
                            year_three = new_row[7:10]
                    else:
                        year_three = new_row[7:10]
                else:
                    year_two = new_row[4:6]
                    if academic_year_3:
                        if academic_year_3['term_3_start_date']:
                            year_three = new_row[6:9]
                        else:
                            year_three = new_row[6:8]
                    else:
                        year_three = new_row[6:8]
            else:
                year_two = new_row[4:7]
                year_three = new_row[7:10]

        else:
            year_one = new_row[1:3]
            if academic_year_2:
                if academic_year_2['term_3_start_date']:
                    year_two = new_row[3:6]
                    if academic_year_3:
                        if academic_year_3['term_3_start_date']:
                            year_three = new_row[6:9]
                        else:
                            year_three = new_row[6:8]
                    else:
                        year_three = new_row[6:8]
                else:
                    year_two = new_row[3:5]
                    if academic_year_3:
                        if academic_year_3['term_3_start_date']:
                            year_three = new_row[5:8]
                        else:
                            year_three = new_row[5:7]
                    else:
                        year_three = new_row[5:7]
            else:
                year_two = new_row[3:5]
                year_three = new_row[5:7]

        st_subject = Subject.objects.get(name=subject['name'])
        new_counter = 1
        for cell1, cell2, cell3 in zip(year_one, year_two, year_three):
            # Cell One
            cell1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para1 = cell1.paragraphs[0]
            para1.add_run(style=verdana2)
            para1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para1.runs[0].font.color.rgb = RGBColor(46, 139, 87)
            para1.runs[0].font.size = Pt(8)

            # Cell Two
            cell2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            para2 = cell2.paragraphs[0]
            para2.add_run(style=verdana2)
            para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para2.runs[0].font.color.rgb = RGBColor(46, 139, 87)
            para2.runs[0].font.size = Pt(8)

            # Cell Three
            if cell3 is not None:
                cell3.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                para3 = cell3.paragraphs[0]
                para3.add_run(style=verdana2)
                para3.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para3.runs[0].font.color.rgb = RGBColor(46, 139, 87)
                para3.runs[0].font.size = Pt(8)

            try:
                academic_year_one = AcademicYear.objects.get(name=academic_year_names[0])
                results_obj = Result.objects.filter(student=student, subject=st_subject, academic_year=academic_year_one, academic_term=new_counter)
                if results_obj.exists():
                    results = StudentResultsSerializer(results_obj, many=True).data
                    score = float(results[0]['score'])
                    if score >= 80:
                        cell1.paragraphs[0].runs[0].text = 'A1'
                    elif 80 > score >= 75:
                        cell1.paragraphs[0].runs[0].text = 'B2'
                    elif 75 > score >= 70:
                        cell1.paragraphs[0].runs[0].text = 'B3'
                    elif 70 > score >= 65:
                        cell1.paragraphs[0].runs[0].text = 'C4'
                    elif 65 > score >= 60:
                        cell1.paragraphs[0].runs[0].text = 'C5'
                    elif 60 > score >= 55:
                        cell1.paragraphs[0].runs[0].text = 'C6'
                    elif 55 > score >= 50:
                        cell1.paragraphs[0].runs[0].text = 'D7'
                    elif 50 > score >= 40:
                        cell1.paragraphs[0].runs[0].text = 'E8'
                    elif score < 40:
                        cell1.paragraphs[0].runs[0].text = 'F9'
                else:
                    cell1.paragraphs[0].runs[0].text = 'N/A'

            except IndexError:
                cell1.paragraphs[0].runs[0].text = 'N/A'

            try:
                academic_year_two = AcademicYear.objects.get(name=academic_year_names[1])
                results_obj = Result.objects.filter(student=student, subject=st_subject, academic_year=academic_year_two,
                                                    academic_term=new_counter)
                if results_obj.exists():
                    results = StudentResultsSerializer(results_obj, many=True).data
                    score = float(results[0]['score'])
                    if score >= 80:
                        cell2.paragraphs[0].runs[0].text = 'A1'
                    elif 80 > score >= 75:
                        cell2.paragraphs[0].runs[0].text = 'B2'
                    elif 75 > score >= 70:
                        cell2.paragraphs[0].runs[0].text = 'B3'
                    elif 70 > score >= 65:
                        cell2.paragraphs[0].runs[0].text = 'C4'
                    elif 65 > score >= 60:
                        cell2.paragraphs[0].runs[0].text = 'C5'
                    elif 60 > score >= 55:
                        cell2.paragraphs[0].runs[0].text = 'C6'
                    elif 55 > score >= 50:
                        cell2.paragraphs[0].runs[0].text = 'D7'
                    elif 50 > score >= 40:
                        cell2.paragraphs[0].runs[0].text = 'E8'
                    elif score < 40:
                        cell2.paragraphs[0].runs[0].text = 'F9'
                else:
                    cell2.paragraphs[0].runs[0].text = 'N/A'

            except IndexError:
                cell2.paragraphs[0].runs[0].text = 'N/A'

            try:
                academic_year_three = AcademicYear.objects.get(name=academic_year_names[2])
                results_obj = Result.objects.filter(student=student, subject=st_subject, academic_year=academic_year_three,
                                                    academic_term=new_counter)
                if results_obj.exists():
                    results = StudentResultsSerializer(results_obj, many=True).data
                    score = float(results[0]['score'])
                    if cell3 is not None:
                        if score >= 80:
                            cell3.paragraphs[0].runs[0].text = 'A1'
                        elif 80 > score >= 75:
                            cell3.paragraphs[0].runs[0].text = 'B2'
                        elif 75 > score >= 70:
                            cell3.paragraphs[0].runs[0].text = 'B3'
                        elif 70 > score >= 65:
                            cell3.paragraphs[0].runs[0].text = 'C4'
                        elif 65 > score >= 60:
                            cell3.paragraphs[0].runs[0].text = 'C5'
                        elif 60 > score >= 55:
                            cell3.paragraphs[0].runs[0].text = 'C6'
                        elif 55 > score >= 50:
                            cell3.paragraphs[0].runs[0].text = 'D7'
                        elif 50 > score >= 40:
                            cell3.paragraphs[0].runs[0].text = 'E8'
                        elif score < 40:
                            cell3.paragraphs[0].runs[0].text = 'F93'
                else:
                    if cell3 is not None:
                        cell3.paragraphs[0].runs[0].text = 'N/A'

            except IndexError:
                if cell3 is not None:
                    cell3.paragraphs[0].runs[0].text = 'N/A'

            new_counter += 1

    for row in r_table.rows:
        row.height = Inches(0.3)

    divider1_para = doc.add_paragraph("divider")
    # Hint
    hint = doc.add_paragraph("GRADING SYSTEM")
    hint.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    hint.runs[0].font.size = Pt(8)
    hint.runs[0].font.bold = True
    hint.paragraph_format.space_after = Pt(0)
    hint.runs[0].font.underline = True
    hint.runs[0].font.color.rgb = RGBColor(46, 139, 87)
    hint_shd = OxmlElement('w:shd')
    hint_shd.set(qn('w:val'), 'clear')
    hint_shd.set(qn('w:color'), 'auto')
    hint_shd.set(qn('w:fill'), 'FFF6D3')
    hint.paragraph_format.element.get_or_add_pPr().append(hint_shd)
    divider2_para = doc.add_paragraph("divider")

    # Grading System and Signature Table
    grad_sign_table = doc.add_table(2, 4)
    for cell in grad_sign_table.columns[1].cells:
        cell.width = Inches(4)

    for cell in grad_sign_table.columns[2].cells:
        cell.width = Inches(7)
    # Grading Cell
    grad_cell = grad_sign_table.rows[0].cells[2]
    grad_cell_para = grad_cell.paragraphs[0]
    a1 = grad_cell_para.add_run("[ A1 ]  EXCELLENT (100-80)%")
    grad_cell_para.add_run().add_break()
    b2 = grad_cell_para.add_run("[ B2 ]  VERY GOOD (79-75)%")
    grad_cell_para.add_run().add_break()
    b3 = grad_cell_para.add_run("[ B3 ]  GOOD (74-70)%")
    grad_cell_para.add_run().add_break()
    c4 = grad_cell_para.add_run("[ C4 ]  CREDIT (69-65)%")
    grad_cell_para.add_run().add_break()
    c5 = grad_cell_para.add_run("[ C5 ]  CREDIT (64-60)%")
    grad_cell_para.add_run().add_break()
    c6 = grad_cell_para.add_run("[ C6 ]  CREDIT (59-55)%")
    grad_cell_para.add_run().add_break()
    d7 = grad_cell_para.add_run("[ D7 ]  PASS (54-50)%")
    grad_cell_para.add_run().add_break()
    e8 = grad_cell_para.add_run("[ E8 ]  WEAK PASS(49-40)%")
    grad_cell_para.add_run().add_break()
    f9 = grad_cell_para.add_run("[ F9 ]  FAIL (39-0)%")
    grad_cell_para.add_run().add_break()
    dash = grad_cell_para.add_run("[ - ]  NO EXAMINATION WAS CONDUCTED AT THE END OF THE SEMESTER")
    grad_cell_para.add_run().add_break()
    n_a = grad_cell_para.add_run("[ N/A ]  NO RESULT RECORD FOUND")

    sign_cell = grad_sign_table.rows[1].cells[1]
    sign_cell_para = sign_cell.paragraphs[0]
    sign_cell_para.add_run().add_picture(f"{student_data['school']['head_signature'].replace(f'{base_url(request)}', '').replace('/', '', 1)}", height=Inches(0.93))
    sign_cell_para.add_run().add_break()
    sign_name = sign_cell_para.add_run(student_data['school']['head_name'])
    sign_cell_para.add_run().add_break()
    sign_title = sign_cell_para.add_run("(HEADMASTER)")

    grad_sign_items = [a1, b2, b3, c4, c5, c6, d7, e8, f9, dash, n_a, sign_name, sign_title, divider1_para, divider2_para]
    for item in grad_sign_items:
        if grad_sign_items.index(item) == 13:
            item.runs[0].font.size = Pt(10)
            item.paragraph_format.space_after = Pt(0)
            item.runs[0].font.color.rgb = RGBColor(255, 246, 211)
            divider_shd = OxmlElement('w:shd')
            item.paragraph_format.space_after = Pt(0)
            divider_shd.set(qn('w:val'), 'clear')
            divider_shd.set(qn('w:color'), 'auto')
            divider_shd.set(qn('w:fill'), 'FFF6D3')
            item.paragraph_format.element.get_or_add_pPr().append(divider_shd)

        elif grad_sign_items.index(item) == 14:
            item.runs[0].font.size = Pt(10)
            item.paragraph_format.space_after = Pt(5)
            item.runs[0].font.color.rgb = RGBColor(255, 246, 211)
            divider_shd = OxmlElement('w:shd')
            item.paragraph_format.space_after = Pt(0)
            divider_shd.set(qn('w:val'), 'clear')
            divider_shd.set(qn('w:color'), 'auto')
            divider_shd.set(qn('w:fill'), 'FFF6D3')
            item.paragraph_format.element.get_or_add_pPr().append(divider_shd)

        else:
            item.font.size = Pt(8)
            item.font.color.rgb = RGBColor(46, 139, 87)

    for row in grad_sign_table.rows:
        for cell in row.cells:
            cell_xml = cell._tc
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), 'FFF6D3')
            cell_xml.get_or_add_tcPr().append(shd)

    # Bottom Bar
    bottom_bar_paragraph = doc.add_paragraph("Powered By: EduAAP")
    bottom_bar_paragraph.runs[0].font.size = Pt(7.5)
    bottom_bar_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    bottom_bar_paragraph.paragraph_format.space_after = Pt(0)
    bottom_bar_paragraph.runs[0].font.color.rgb = RGBColor(255, 255, 0)
    # Bottom Bar Shade
    bottom_bar_shd = OxmlElement('w:shd')
    bottom_bar_shd.set(qn('w:val'), 'clear')
    bottom_bar_shd.set(qn('w:color'), 'auto')
    bottom_bar_shd.set(qn('w:fill'), '2E8B57')
    bottom_bar_paragraph.paragraph_format.element.get_or_add_pPr().append(bottom_bar_shd)

    # Save The Doc File
    doc_byte = io.BytesIO()
    doc.save(doc_byte)

    if settings.DEBUG:
        file_path = f"{get_school_folder(student_data['school']['name'])}/students/{student_data['user']['username']}/{student_data['user']['first_name']}_{student_data['user']['last_name']}_transcript.docx"
        if default_storage.exists(file_path):
            default_storage.delete(file_path)

        save_file = default_storage.save(file_path, doc_byte)

        return f"http://localhost:8000{default_storage.url(save_file)}"

    else:
        file_path = f"media/{get_school_folder(student_data['school']['name'])}/students/{student_data['user']['username']}/{student_data['user']['first_name']}_{student_data['user']['last_name']}_transcript.docx"
        if default_storage.exists(file_path):
            default_storage.delete(file_path)

        save_file = default_storage.save(file_path, doc_byte)

        return default_storage.url(save_file)


# Get current academic year
def get_current_academic_year(school_user, user_data):
    academic_years = AcademicYear.objects.filter(school=school_user.school).order_by('-start_date')
    current_date = timezone.now().date()
    if academic_years.exists():
        current_year = academic_years[0]
        if current_year.start_date <= current_date:
            if not current_date > current_year.term_1_end_date:
                user_data['current_term'] = 1

            elif not current_date > current_year.term_2_end_date:
                user_data['current_term'] = 2

            elif not current_year.term_3_start_date:
                user_data['current_term'] = 2

            elif current_year.term_3_start_date:
                user_data['current_term'] = 3

            user_data['academic_year'] = AcademicYearSerializer(current_year).data

        elif academic_years.count() > 1:
            previous_year = academic_years[1]
            if not current_date > previous_year.term_1_end_date:
                user_data['current_term'] = 1

            elif not current_date > previous_year.term_2_end_date:
                user_data['current_term'] = 2

            elif not previous_year.term_3_start_date:
                user_data['current_term'] = 2

            elif previous_year.term_3_start_date:
                user_data['current_term'] = 3

            user_data['academic_year'] = AcademicYearSerializer(previous_year).data

        else:
            return Response({'ms': 'No current academic year, contact you school administrator'}, status=201)

    else:
        return Response({'ms': 'No current academic year, contact you school administrator'}, status=201)


# Hod subject Assignments Data
def get_hod_subject_assignments(hod, academic_year):
    classes = []
    staff = []
    subjects = []
    term_one_subject_assignments = SubjectAssignmentSerializer(
        SubjectAssignment.objects.filter(hod=hod, academic_year=academic_year, academic_term=1), many=True).data
    term_two_subject_assignments = SubjectAssignmentSerializer(
        SubjectAssignment.objects.filter(hod=hod, academic_year=academic_year, academic_term=2), many=True).data
    term_three_subject_assignments = SubjectAssignmentSerializer(
        SubjectAssignment.objects.filter(hod=hod, academic_year=academic_year, academic_term=3), many=True).data

    try:
        department = DepartmentSerializer(Department.objects.get(hod=hod)).data
        classes_data = ClasseSerializer(Classe.objects.filter(school=hod.school, is_active=True), many=True).data

        for item in department['teachers']:
            staff.append({
                'name': f"{item['user']['first_name']} {item['user']['last_name']}",
                'stf_id': item['staff_id'],
                'img': item['img'],
                'subjects': item['subjects'],
                'role': item['role']
            })

        for item in classes_data:
            classes.append(f"{item['name']} FORM {item['students_year']}")

        for item in department['subjects']:
            subjects.append(item['name'])

    except Department.DoesNotExist:
        pass

    data = {
        'term_one': term_one_subject_assignments,
        'term_two': term_two_subject_assignments,
        'term_three': term_three_subject_assignments,
        'subjects': subjects,
        'staff': staff,
        'classes': classes,
    }

    return data


# Teacher upload results
def teacher_results_upload(staff, year, term):
    subject_assignments = SubjectAssignment.objects.select_related('students_class', 'subject').filter(
        school=staff.school,
        teacher=staff,
        academic_year=year,
        academic_term=term,
    )
    results = Result.objects.select_related('student', 'subject').filter(
        school=staff.school,
        teacher=staff,
        academic_year=year,
        academic_term=term,
    )
    serialized_results = ResultsSerializer(results, many=True).data
    teacher_students_without_results = []
    teacher_students_with_results = []
    serialized_subject_assignments = SubjectAssignmentSerializer(subject_assignments, many=True).data
    if serialized_subject_assignments:
        for assign_item in serialized_subject_assignments:
            subject_assignment_students = assign_item['students_class']['students']
            subject_assignment_subject_name = assign_item['subject']['name']
            subject_assignment_st_year = assign_item['students_class']['students_year']
            subject_assignment_class_name = assign_item['students_class']['name']
            students_without_results = []
            students_with_results = []

            for st_item in subject_assignment_students:
                subject_assignment_st_name = f"{st_item['user']['first_name']} {st_item['user']['last_name']}"
                subject_assignment_st_id = st_item['st_id']
                students_without_results.append({'name': subject_assignment_st_name, 'st_id': subject_assignment_st_id})
                if serialized_results:
                    for result_item in serialized_results:
                        result_st_id = result_item['student']['st_id']
                        results_subject_name = result_item['subject']['name']
                        if results_subject_name == subject_assignment_subject_name and subject_assignment_st_id == result_st_id:
                            students_without_results = [st for st in students_without_results if
                                                        st['st_id'] != result_st_id]
                            students_with_results.append(
                                {'student': result_item['student'], 'score': result_item['score']})

            teacher_students_without_results.append({
                'subject': subject_assignment_subject_name,
                'class_name': subject_assignment_class_name,
                'students_year': subject_assignment_st_year,
                'students': students_without_results
            })

            students_with_results = sorted(students_with_results, key=lambda x: float(x['score']), reverse=True)
            teacher_students_with_results.append({
                'subject': subject_assignment_subject_name,
                'class_name': subject_assignment_class_name,
                'students_year': subject_assignment_st_year,
                'students': students_with_results
            })

        return [teacher_students_without_results, teacher_students_with_results]


# Notifications messages
def get_staff_messages(request_user):
    try:
        staff = request_user.staff
        messages = []
        messages_sent = StaffNotificationSerializer(StaffNotification.objects.filter(sent_by_hod=staff), many=True).data
        if messages_sent:
            for msg in messages_sent:
                messages.append(msg)

        messages_received = StaffNotificationSerializer(StaffNotification.objects.filter(send_to=staff), many=True).data
        if messages_received:
            for msg in messages_received:
                messages.append(msg)

        messages = sorted(messages, key=lambda x: x['date_time'], reverse=True)
        response_data = []
        for msg in messages:
            msg['date_time'] = format_relative_date_time(msg['date_time'], True, True)
            response_data.append(msg)

        return response_data

    except User.staff.RelatedObjectDoesNotExist:
        head = request_user.head
        messages = []
        messages_sent = StaffNotificationSerializer(StaffNotification.objects.filter(sent_by_head=head), many=True).data
        if messages_sent:
            for msg in messages_sent:
                messages.append(msg)

        messages = sorted(messages, key=lambda x: x['date_time'], reverse=True)

        response_data = []
        for msg in messages:
            msg['date_time'] = format_relative_date_time(msg['date_time'], True, True)
            response_data.append(msg)

        return response_data


# Convert datetime
def format_relative_date_time(utc_date, day_name, time):
    # Convert the UTC string to a datetime object
    utc_datetime = datetime.strptime(utc_date, "%Y-%m-%dT%H:%M:%S.%fZ")

    # Get the current UTC datetime
    current_utc_datetime = datetime.utcnow()

    # Check if it's today
    if utc_datetime.date() == current_utc_datetime.date():
        if day_name and time:
            return f"Today, {utc_datetime.strftime('%B %d, %Y at %H:%M')}"
        elif day_name and not time:
            return f"Today, {utc_datetime.strftime('%B %d, %Y')}"
        elif not day_name and time:
            return utc_datetime.strftime('%B %d, %Y at %H:%M')
        elif not day_name and not time:
            return utc_datetime.strftime('%B %d, %Y')

    elif utc_datetime.date() == (current_utc_datetime - timedelta(days=1)).date():
        if day_name and time:
            return f"Yesterday, {utc_datetime.strftime('%B %d, %Y at %H:%M')}"
        elif day_name and not time:
            return f"Yesterday, {utc_datetime.strftime('%B %d, %Y')}"
        elif not day_name and time:
            return utc_datetime.strftime('%B %d, %Y at %H:%M')
        elif not day_name and not time:
            return utc_datetime.strftime('%B %d, %Y')

    else:
        if day_name and time:
            return utc_datetime.strftime('%A, %B %d, %Y at %H:%M')
        elif day_name and not time:
            return utc_datetime.strftime('%A, %B %d, %Y')
        elif not day_name and time:
            return utc_datetime.strftime('%B %d, %Y at %H:%M')
        elif not day_name and not time:
            return utc_datetime.strftime('%B %d, %Y')

